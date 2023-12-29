"""Wraps paragraphs formatted with the OneNote code style font in a Markdown code block"""
import docx
import argparse

CODE_STYLE_FONT_NAME: str = "Consolas"
SPACE_UNICODE_CODE: int = 0x20
NON_BREAKING_SPACE_UNICODE_CODE: int = 0xA0


def replace_leading_spaces(text: str) -> str:
    """Replaces leading spaces with non-breaking spaces plus a space.

    Substituting spaces with non-breaking spaces is necessary for Pandoc to
    output code blocks with the correct indentation.

    Parameters
    ----------
    text : str
        The text to update.

    Returns
    -------
    str
        The updated text.
    """
    if len(text) == 0:
        return ""

    new_text = ""
    prev_char = text[0]
    leading = True

    for c in text:
        # If we encounter a non-whitespace character, then we have reached the
        # beginning of the text and no longer need to replace spaces
        if ord(c) != SPACE_UNICODE_CODE and ord(c) != NON_BREAKING_SPACE_UNICODE_CODE:
            leading = False

        if (
            leading
            and ord(c) == SPACE_UNICODE_CODE
            and ord(prev_char) != NON_BREAKING_SPACE_UNICODE_CODE
        ):
            # Substitute a space with a non-breaking space plus a normal space
            new_text += chr(NON_BREAKING_SPACE_UNICODE_CODE) + chr(SPACE_UNICODE_CODE)
        else:
            # Otherwise, just append the character as normal
            new_text += c

        prev_char = c

    return new_text


def close_code_block(paragraph: docx.text.paragraph.Paragraph) -> None:
    """Appends triple backticks to the end of a paragraph.

    Parameters
    ----------
    paragraph : docx.text.paragraph.Paragraph
        The docx paragraph to update.
    """
    index = 1
    # Get the last non-empty run inserted into the paragraph and insert three
    # backticks to signal the end of the block
    last_run = paragraph.runs[len(paragraph.runs) - index]
    while len(last_run.text.strip()) == 0 and index <= (len(paragraph.runs) - 1):
        index += 1
        last_run = paragraph.runs[len(paragraph.runs) - index]

    # The run text already contains a line break so a newline does not
    # need to be added
    last_run.text = last_run.text + "```"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("filename")
    args = parser.parse_args()

    doc = docx.Document(args.filename)

    in_code_block = False
    inserted_paragraph = None

    for p in doc.paragraphs:
        if len(p.runs) == 0:
            continue

        font = p.runs[0].font.name

        if (
            font == CODE_STYLE_FONT_NAME
            and len(p.text.strip()) > 0
            and not in_code_block
        ):
            # Generate a new paragraph before the beginning of the code block
            inserted_paragraph = p.insert_paragraph_before()
            # Append three backticks to the current paragraph text to signal the
            # start of the block
            p.text = "```\n" + p.text
            in_code_block = True

        if font == CODE_STYLE_FONT_NAME and in_code_block:
            run = inserted_paragraph.add_run(replace_leading_spaces(p.text.rstrip()))
            # We must use line breaks instead of paragraph breaks otherwise
            # pandoc will output incorrect newlines into the code block
            run.add_break()

            # Remove the old paragraph. The standard paragraph cannot be used
            # because it contains a paragraph break at the end which will cause
            # pandoc to output empty newlines
            try:
                xml = p._element
                xml.getparent().remove(xml)
                xml._p = xml._element = None
            except AttributeError:
                continue

        if font != CODE_STYLE_FONT_NAME and in_code_block:
            close_code_block(inserted_paragraph)
            in_code_block = False

    # If a code block is at the end of document, close it now
    if in_code_block:
        close_code_block(inserted_paragraph)

    doc.save(args.filename)


if __name__ == "__main__":
    main()
